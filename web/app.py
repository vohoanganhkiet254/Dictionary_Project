# -*- coding: utf-8 -*-
"""
Web app kiểm tra & sửa chính tả OCR cho file .docx.

Chạy:
    python corpus/web/app.py
rồi mở http://127.0.0.1:5000

Chức năng:
    - Tải .docx lên -> hiển thị nội dung.
    - Từ sai chính tả được gạch chân; bấm vào -> hiện gợi ý sửa.
    - Sửa trực tiếp trên web ("Kiểm tra lại" để soát lại).
    - Tải về .docx đã chỉnh sửa.
"""
import os
import sys
import io

from flask import Flask, request, jsonify, send_file, render_template

import docx

# --- nạp bộ kiểm tra chính tả từ corpus/scripts ---------------------------- #
HERE = os.path.dirname(os.path.abspath(__file__))
SCRIPTS = os.path.normpath(os.path.join(HERE, "..", "scripts"))
sys.path.insert(0, SCRIPTS)
from spellcheck_ocr import Dictionary, SpellChecker  # noqa: E402

app = Flask(__name__)
app.config["MAX_CONTENT_LENGTH"] = 25 * 1024 * 1024   # giới hạn 25MB

_dic = Dictionary()
_checker = SpellChecker(_dic)


def check_paragraph(text: str):
    """Trả danh sách lỗi của 1 đoạn: vị trí + từ gốc + gợi ý."""
    issues = []
    for it in _checker.check(text):
        issues.append({
            "start": it.start,
            "end": it.end,
            "original": it.original,
            "type": it.err_type,
            "suggestions": [s for s, _t, _sc in it.suggestions],
        })
    return issues


@app.route("/")
def index():
    return render_template("index.html")


@app.route("/upload", methods=["POST"])
def upload():
    f = request.files.get("file")
    if not f or not f.filename.lower().endswith(".docx"):
        return jsonify({"error": "Vui lòng chọn file .docx"}), 400
    try:
        document = docx.Document(io.BytesIO(f.read()))
    except Exception as e:
        return jsonify({"error": f"Không đọc được file: {e}"}), 400
    paras = [{"text": p.text, "issues": check_paragraph(p.text)}
             for p in document.paragraphs]
    return jsonify({"filename": f.filename, "paragraphs": paras})


@app.route("/check", methods=["POST"])
def check():
    data = request.get_json(force=True)
    paras = data.get("paragraphs", [])
    out = [{"text": t, "issues": check_paragraph(t)} for t in paras]
    return jsonify({"paragraphs": out})


@app.route("/download", methods=["POST"])
def download():
    data = request.get_json(force=True)
    paras = data.get("paragraphs", [])
    document = docx.Document()
    for t in paras:
        document.add_paragraph(t)
    buf = io.BytesIO()
    document.save(buf)
    buf.seek(0)
    return send_file(
        buf, as_attachment=True, download_name="da-chinh-sua.docx",
        mimetype="application/vnd.openxmlformats-officedocument.wordprocessingml.document",
    )


if __name__ == "__main__":
    app.run(host="127.0.0.1", port=5000, debug=True)
